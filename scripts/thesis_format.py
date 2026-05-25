#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
论文格式学习与应用工具 v2
从参考文档学习格式规则，然后批量应用到其他文档。

用法:
  python thesis_format.py learn reference.docx [rules.json]
  python thesis_format.py apply input.docx rules.json [output.docx]
  python thesis_format.py check input.docx rules.json
"""
import json
import re
import sys
from collections import Counter
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================
# 段落分类
# ============================================================

def classify_paragraph(para):
    """根据段落内容和样式判断类型。"""
    text = para.text.strip()
    if not text:
        return "empty"

    style_name = (para.style.name if para.style else "").lower()

    # TOC styles
    if style_name.startswith("toc"):
        return "toc"

    # Heading styles (Word 内置样式)
    if style_name == "heading 1":
        return "heading1"
    if style_name == "heading 2":
        return "heading2"
    if style_name == "heading 3":
        return "heading3"
    if style_name.startswith("heading"):
        return "heading_other"

    # 编号标题检测（自动编号，Normal 样式）
    # H1: "1 标题" "2标题"（一个或多个空格/全角空格）
    if re.match(r'^\d+[\s　]+\S', text) and not re.match(r'^\d+\.\d', text):
        if not re.match(r'^\[\d+\]', text) and len(text) < 80:
            return "heading1"
    # H2: "1.1 标题" "1.1标题" "1.1  标题"
    if re.match(r'^\d+\.\d+[\s　]*\S', text) and not re.match(r'^\d+\.\d+\.\d', text):
        return "heading2"
    # H3: "1.1.1 标题" "1.1.1标题" "1.1.1  标题"
    if re.match(r'^\d+\.\d+\.\d+[\s　]*\S', text) and not re.match(r'^\d+\.\d+\.\d+\.\d', text):
        return "heading3"
    # H4: "1.1.1.1 标题"
    if re.match(r'^\d+\.\d+\.\d+\.\d+[\s　]*\S', text):
        return "heading4"

    # TOC title
    if text.replace(" ", "").replace("　", "") in ("目录", "目　录"):
        return "toc_title"

    # 论文标题：居中 + (大字号 或 Normal样式 + 无编号 + 短文本)
    align = para.alignment
    is_centered = (align == WD_ALIGN_PARAGRAPH.CENTER or str(align) == "CENTER (1)")
    if is_centered and style_name in ("normal", ""):
        # 检查 run 级字号
        run_large = False
        if para.runs and para.runs[0].font.size:
            run_large = para.runs[0].font.size.pt >= 14
        # 检查样式级字号
        style_large = False
        style_info = get_style_info(para)
        if style_info.get("style_size_pt", 0) >= 14:
            style_large = True
        # 居中 + 大字号 + 不以数字开头 → 标题
        if (run_large or style_large) and not re.match(r'^\d', text):
            return "title"
        # 居中 + Normal样式 + 不以数字开头 + 不是目录/摘要等 → 也可能是标题
        if not re.match(r'^\d', text) and len(text) < 50 and not any(
            text.startswith(k) for k in ("目", "摘", "Abstract", "Key", "参", "致", "附", "目 录")
        ):
            return "title"

    # Reference heading
    if text.replace(" ", "") in ("参考文献",):
        return "ref_heading"

    # Acknowledgment heading
    if text.replace(" ", "").replace("　", "") in ("致谢",):
        return "ack_heading"

    # References
    if re.match(r'^\[\d+\]', text):
        return "reference"

    # Figure/table captions
    if re.match(r'^(图|表|Fig|Table)\s*\d', text):
        return "caption"

    # 续表 captions
    if text.startswith("续表"):
        return "caption"

    # Table notes
    if text.startswith("注：") or text.startswith("注:"):
        return "table_note"

    # Abstract labels
    if re.match(r'^摘\s*要[：:]', text):
        return "abstract_cn"
    if re.match(r'^关键\s*词[：:]', text):
        return "keywords_cn"
    if re.match(r'^Abstract[：:\s]', text, re.IGNORECASE):
        return "abstract_en"
    if re.match(r'^Key\s*words[：:\s]', text, re.IGNORECASE):
        return "keywords_en"

    return "body"


def is_front_matter(para, seen_toc):
    """判断段落是否属于封面/声明等前置部分。"""
    if seen_toc:
        return False
    ptype = classify_paragraph(para)
    if ptype in ("toc_title", "toc"):
        return False
    return True


# ============================================================
# 格式提取
# ============================================================

def get_font_info(para):
    """从段落的所有 run 提取字体信息（多数投票）。"""
    if not para.runs:
        return {}

    fonts_en = []
    fonts_cn = []
    sizes = []
    bolds = []

    from lxml import etree
    ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    for run in para.runs:
        text = run.text.strip()
        if not text:
            continue
        f = run.font
        if f.name:
            fonts_en.append(f.name)
        if f.size:
            sizes.append(round(f.size.pt, 1))
        bolds.append(bool(f.bold))

        # 中文字体
        rpr = run._element.find(f"{ns_w}rPr")
        if rpr is not None:
            fonts_elem = rpr.find(f"{ns_w}rFonts")
            if fonts_elem is not None:
                east = fonts_elem.get(f"{ns_w}eastAsia")
                if east:
                    fonts_cn.append(east)

    info = {}
    if fonts_cn:
        info["font_cn"] = Counter(fonts_cn).most_common(1)[0][0]
    if fonts_en:
        info["font_en"] = Counter(fonts_en).most_common(1)[0][0]
    if sizes:
        info["size_pt"] = Counter(sizes).most_common(1)[0][0]
    if bolds:
        majority = sum(bolds) > len(bolds) / 2
        if majority:
            info["bold"] = True
        else:
            info["bold"] = False

    return info


def get_para_info(para):
    """提取段落格式信息。"""
    info = {}
    pf = para.paragraph_format

    if pf.alignment is not None:
        info["alignment"] = str(pf.alignment)

    if pf.line_spacing is not None:
        if isinstance(pf.line_spacing, float):
            info["line_spacing_multiple"] = pf.line_spacing
        else:
            info["line_spacing_pt"] = round(pf.line_spacing.pt, 1)

    if pf.line_spacing_rule is not None:
        info["line_spacing_rule"] = str(pf.line_spacing_rule)

    if pf.space_before is not None:
        info["space_before_pt"] = round(pf.space_before.pt, 1)
    if pf.space_after is not None:
        info["space_after_pt"] = round(pf.space_after.pt, 1)

    if pf.first_line_indent is not None:
        info["first_line_indent_pt"] = round(pf.first_line_indent.pt, 1)

    if pf.left_indent is not None:
        info["left_indent_pt"] = round(pf.left_indent.pt, 1)

    return info


def get_style_info(para):
    """从段落的 Word 样式提取格式信息。"""
    from lxml import etree
    ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    style = para.style
    if not style:
        return {}

    info = {}

    # 样式的字体
    style_elem = style._element
    rpr = style_elem.find(f"{ns_w}rPr")
    if rpr is not None:
        fonts = rpr.find(f"{ns_w}rFonts")
        if fonts is not None:
            east = fonts.get(f"{ns_w}eastAsia")
            if east:
                info["style_font_cn"] = east
            ascii_f = fonts.get(f"{ns_w}ascii")
            if ascii_f:
                info["style_font_en"] = ascii_f
        sz = rpr.find(f"{ns_w}sz")
        if sz is not None:
            val = sz.get(f"{ns_w}val")
            if val:
                info["style_size_pt"] = int(val) / 2
        b = rpr.find(f"{ns_w}b")
        if b is not None:
            info["style_bold"] = True

    # 样式的段落格式
    ppr = style_elem.find(f"{ns_w}pPr")
    if ppr is not None:
        jc = ppr.find(f"{ns_w}jc")
        if jc is not None:
            val = jc.get(f"{ns_w}val")
            if val:
                info["style_alignment"] = val

    return info


# ============================================================
# 学习格式
# ============================================================

def learn_format(docx_path):
    """从参考文档学习格式规则。"""
    doc = Document(str(docx_path))
    rules = {}

    # 页面设置
    section = doc.sections[0]
    rules["_page"] = {
        "top_cm": round(section.top_margin.cm, 2),
        "bottom_cm": round(section.bottom_margin.cm, 2),
        "left_cm": round(section.left_margin.cm, 2),
        "right_cm": round(section.right_margin.cm, 2),
        "header_cm": round(section.header_distance.cm, 2),
        "footer_cm": round(section.footer_distance.cm, 2),
    }

    # 跳过封面/声明，从 TOC 之后开始收集
    seen_toc = False
    body_started = False
    samples = {}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        ptype = classify_paragraph(para)

        # 标记 TOC，但仍然学习 toc_title 格式
        if ptype == "toc":
            seen_toc = True
            continue
        if ptype == "toc_title":
            seen_toc = True
            # 继续学习 toc_title 格式

        # 跳过封面/声明
        if not seen_toc:
            continue

        # 等到第一个标题出现才开始收集
        if not body_started:
            if ptype in ("heading1", "heading2", "heading3", "ref_heading", "ack_heading", "toc_title"):
                body_started = True
            else:
                continue

        if ptype not in samples:
            samples[ptype] = []
        samples[ptype].append({
            "text": text[:50],
            "font": get_font_info(para),
            "para": get_para_info(para),
            "style": get_style_info(para),
        })

    # 每种类型取最常见的格式
    for ptype, items in samples.items():
        if not items:
            continue

        merged = {}

        # 合并 run 级字体信息（多数投票）
        for key in ("font_cn", "font_en", "size_pt", "bold"):
            values = [item["font"].get(key) for item in items if item["font"].get(key) is not None]
            if values:
                if key == "size_pt":
                    # 取最常见的字号
                    merged[key] = Counter(values).most_common(1)[0][0]
                elif key == "bold":
                    merged[key] = sum(values) > len(values) / 2
                else:
                    merged[key] = Counter(values).most_common(1)[0][0]

        # 合并样式级信息（作为 fallback）
        for key in ("style_font_cn", "style_font_en", "style_size_pt", "style_bold", "style_alignment"):
            values = [item["style"].get(key) for item in items if item["style"].get(key) is not None]
            if values:
                merged[key] = Counter(values).most_common(1)[0][0]

        # 合并段落格式（取最常见的）
        for key in ("alignment", "line_spacing_pt", "line_spacing_rule", "line_spacing_multiple",
                     "space_before_pt", "space_after_pt", "first_line_indent_pt", "left_indent_pt"):
            values = [item["para"].get(key) for item in items if item["para"].get(key) is not None]
            if values:
                if isinstance(values[0], (int, float)):
                    merged[key] = Counter(values).most_common(1)[0][0]
                else:
                    merged[key] = Counter(values).most_common(1)[0][0]

        # 只保留有内容的字段
        rules[ptype] = {k: v for k, v in merged.items() if v is not None}

    return rules


# ============================================================
# 应用格式
# ============================================================

ALIGN_MAP = {
    "LEFT (0)": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER (1)": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT (2)": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY (3)": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "DISTRIBUTE (4)": WD_ALIGN_PARAGRAPH.DISTRIBUTE,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def apply_font(run, rule):
    """应用字体规则到 run。"""
    if not rule:
        return
    font = run.font
    if "size_pt" in rule and rule["size_pt"]:
        font.size = Pt(rule["size_pt"])
    if "bold" in rule and rule["bold"] is not None:
        font.bold = rule["bold"]
    if "italic" in rule and rule["italic"] is not None:
        font.italic = rule["italic"]
    if "font_en" in rule and rule["font_en"]:
        font.name = rule["font_en"]

    # 中文字体
    if "font_cn" in rule and rule["font_cn"]:
        from lxml import etree
        ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
        rpr = run._element.find(f"{ns_w}rPr")
        if rpr is None:
            rpr = etree.SubElement(run._element, f"{ns_w}rPr")
        fonts = rpr.find(f"{ns_w}rFonts")
        if fonts is None:
            fonts = etree.SubElement(rpr, f"{ns_w}rFonts")
        fonts.set(f"{ns_w}eastAsia", rule["font_cn"])


def apply_style_fallback(para, rule):
    """当 run 级没有信息时，用样式级信息更新样式定义。"""
    if not rule:
        return
    from lxml import etree
    ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    style = para.style
    if not style:
        return
    style_elem = style._element

    # 更新样式的字体
    if "style_font_cn" in rule or "style_font_en" in rule or "style_size_pt" in rule or "style_bold" in rule:
        rpr = style_elem.find(f"{ns_w}rPr")
        if rpr is None:
            rpr = etree.SubElement(style_elem, f"{ns_w}rPr")
        if "style_font_cn" in rule:
            fonts = rpr.find(f"{ns_w}rFonts")
            if fonts is None:
                fonts = etree.SubElement(rpr, f"{ns_w}rFonts")
            fonts.set(f"{ns_w}eastAsia", rule["style_font_cn"])
        if "style_font_en" in rule:
            fonts = rpr.find(f"{ns_w}rFonts")
            if fonts is None:
                fonts = etree.SubElement(rpr, f"{ns_w}rFonts")
            fonts.set(f"{ns_w}ascii", rule["style_font_en"])
        if "style_size_pt" in rule:
            sz = rpr.find(f"{ns_w}sz")
            if sz is None:
                sz = etree.SubElement(rpr, f"{ns_w}sz")
            sz.set(f"{ns_w}val", str(int(rule["style_size_pt"] * 2)))
        if "style_bold" in rule:
            b = rpr.find(f"{ns_w}b")
            if b is None and rule["style_bold"]:
                etree.SubElement(rpr, f"{ns_w}b")


def apply_para_format(para, rule):
    """应用段落格式规则。"""
    if not rule:
        return
    pf = para.paragraph_format

    if "alignment" in rule and rule["alignment"]:
        pf.alignment = ALIGN_MAP.get(rule["alignment"])

    if "line_spacing_pt" in rule and rule["line_spacing_pt"]:
        pf.line_spacing = Pt(rule["line_spacing_pt"])

    if "space_before_pt" in rule and rule["space_before_pt"] is not None:
        pf.space_before = Pt(rule["space_before_pt"])
    if "space_after_pt" in rule and rule["space_after_pt"] is not None:
        pf.space_after = Pt(rule["space_after_pt"])

    if "first_line_indent_pt" in rule and rule["first_line_indent_pt"] is not None:
        pf.first_line_indent = Pt(rule["first_line_indent_pt"])
    if "left_indent_pt" in rule and rule["left_indent_pt"] is not None:
        pf.left_indent = Pt(rule["left_indent_pt"])


def fix_toc_styles(doc):
    """修复 TOC 样式：宋体 + Times New Roman + 五号(10.5pt)。"""
    from lxml import etree
    ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    styles_elem = doc.styles._element
    for style in styles_elem.findall(f"{ns_w}style"):
        style_id = style.get(f"{ns_w}styleId", "")
        if not style_id.startswith("toc"):
            continue

        # 修复字体
        rpr = style.find(f"{ns_w}rPr")
        if rpr is None:
            rpr = etree.SubElement(style, f"{ns_w}rPr")
        fonts = rpr.find(f"{ns_w}rFonts")
        if fonts is None:
            fonts = etree.SubElement(rpr, f"{ns_w}rFonts")
        fonts.set(f"{ns_w}eastAsia", "宋体")
        fonts.set(f"{ns_w}ascii", "Times New Roman")
        fonts.set(f"{ns_w}hAnsi", "Times New Roman")

        # 修复字号
        sz = rpr.find(f"{ns_w}sz")
        if sz is None:
            sz = etree.SubElement(rpr, f"{ns_w}sz")
        sz.set(f"{ns_w}val", "21")  # 五号 = 10.5pt = 21 half-pt

        # 去掉加粗
        b = rpr.find(f"{ns_w}b")
        if b is not None:
            rpr.remove(b)


def get_official_rule(ptype):
    """返回浙江农林大学官方格式规则。"""
    rules = {
        "heading1": {
            "font_cn": "楷体", "font_en": "Times New Roman",
            "size_pt": 14.0, "bold": True,
            "style_font_cn": "楷体", "style_font_en": "Times New Roman",
            "style_size_pt": 14.0, "style_bold": True,
            "alignment": "CENTER (1)",
            "space_before_pt": 6, "space_after_pt": 6,
            "line_spacing_pt": 20.0,
            "first_line_indent_pt": 0, "left_indent_pt": 0,
        },
        "heading2": {
            "font_cn": "黑体", "font_en": "Times New Roman",
            "size_pt": 12.0, "bold": True,
            "style_font_cn": "黑体", "style_font_en": "Times New Roman",
            "style_size_pt": 12.0, "style_bold": True,
            "space_before_pt": 3, "space_after_pt": 3,
            "line_spacing_pt": 20.0,
            "first_line_indent_pt": 0, "left_indent_pt": 0,
        },
        "heading3": {
            "font_cn": "黑体", "font_en": "Times New Roman",
            "size_pt": 10.5, "bold": False,
            "style_font_cn": "黑体", "style_font_en": "Times New Roman",
            "style_size_pt": 10.5, "style_bold": False,
            "space_before_pt": 3, "space_after_pt": 3,
            "line_spacing_pt": 20.0,
            "first_line_indent_pt": 0, "left_indent_pt": 0,
        },
        "heading4": {
            "font_cn": "宋体", "font_en": "Times New Roman",
            "size_pt": 10.5, "bold": False,
            "style_font_cn": "宋体", "style_font_en": "Times New Roman",
            "style_size_pt": 10.5, "style_bold": False,
            "space_before_pt": 3, "space_after_pt": 3,
            "line_spacing_pt": 20.0,
            "first_line_indent_pt": 0, "left_indent_pt": 0,
        },
    }
    return rules.get(ptype, {})


def renumber_caption(text, chapter_num, fig_count, table_count):
    """重新编号图/表标题。格式：图 X-Y 或 表 X-Y。"""
    # 图：匹配 "图 1-1"、"图1-1"、"图 1‑1"（全角连字符）等
    m = re.match(r'^(图|Fig\.?)\s*(\d+)\s*[-–—‑]\s*(\d+)(.*)', text, re.IGNORECASE)
    if m:
        prefix = m.group(1)
        new_num = f"{chapter_num}-{fig_count + 1}"
        rest = m.group(4)
        return f"{prefix} {new_num}{rest}"

    # 续表：不改编号，只保留
    if text.startswith("续表"):
        return text

    # 表：匹配 "表 1-1"、"表1-1" 等
    m = re.match(r'^(表|Table)\s*(\d+)\s*[-–—‑]\s*(\d+)(.*)', text, re.IGNORECASE)
    if m:
        prefix = m.group(1)
        new_num = f"{chapter_num}-{table_count + 1}"
        rest = m.group(4)
        return f"{prefix} {new_num}{rest}"

    return text


def update_caption_text(para, new_text):
    """更新段落文本（保留第一个 run 的格式）。"""
    if not para.runs:
        return
    # 清空所有 run 的文本
    for run in para.runs:
        run.text = ""
    # 把新文本写入第一个 run
    para.runs[0].text = new_text


def italicize_statistics(para):
    """将 p<0.05, P<0.01, p>0.05 等统计标记设为斜体。"""
    for run in para.runs:
        text = run.text
        # 匹配 p<0.05, P<0.01, p>0.05, p≤0.05, p≥0.05 等
        if re.search(r'[pP]\s*[<>≤≥]\s*0\.\d+', text):
            run.font.italic = True


def update_fields(doc):
    """设置 TOC 域自动更新，但不强制刷新所有域（避免损坏的域报错）。"""
    from lxml import etree
    ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    # 不设置全局 updateFields，只标记 TOC 需要更新
    # Word 打开时用户可以手动 Ctrl+A → F9 刷新
    pass


def check_references(doc):
    """检测文中引用的图/表编号与实际编号是否匹配。"""
    # 收集所有图/表的实际编号
    actual_figs = {}
    actual_tables = {}
    chapter_num = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name if para.style else "").lower()
        if style == "heading 1" or re.match(r'^\d+[\s　]+\S', text):
            if not re.match(r'^\d+\.\d', text):
                chapter_num += 1

        m = re.match(r'^图\s*(\d+[-–—]\d+)', text)
        if m:
            actual_figs[m.group(1)] = text[:30]
        m = re.match(r'^表\s*(\d+[-–—]\d+)', text)
        if m:
            actual_tables[m.group(1)] = text[:30]

    # 检测文中引用
    issues = []
    for para in doc.paragraphs:
        text = para.text
        # 匹配 "图 X-Y"、"表 X-Y" 引用（不在行首的）
        for m in re.finditer(r'(?<!^)(图|表)\s*(\d+\s*[-–—]\s*\d+)', text):
            label = m.group(1)
            num = re.sub(r'\s', '', m.group(2))  # 去空格
            if label == "图" and num not in actual_figs:
                issues.append(f"引用 图{num} 不存在（{text[:40]}）")
            elif label == "表" and num not in actual_tables:
                issues.append(f"引用 表{num} 不存在（{text[:40]}）")

    return issues


def apply_page_setup(doc, page_rule):
    """应用页面设置。"""
    if not page_rule:
        return
    for section in doc.sections:
        section.top_margin = Cm(page_rule.get("top_cm", 2.7))
        section.bottom_margin = Cm(page_rule.get("bottom_cm", 2.7))
        section.left_margin = Cm(page_rule.get("left_cm", 2.7))
        section.right_margin = Cm(page_rule.get("right_cm", 2.7))
        if "header_cm" in page_rule:
            section.header_distance = Cm(page_rule["header_cm"])
        if "footer_cm" in page_rule:
            section.footer_distance = Cm(page_rule["footer_cm"])


def apply_split_font(para, spec):
    """对摘要/关键词段落，标签和内容使用不同字体。

    spec 示例:
    {
        "label_font_cn": "黑体", "label_font_en": "Times New Roman",
        "label_bold": True, "label_size": 10.5,
        "content_font_cn": "楷体", "content_font_en": "Times New Roman",
        "content_bold": False, "content_size": 10.5,
    }
    会自动识别 "摘要：" / "Abstract：" 等标签边界。
    """
    from lxml import etree
    ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    full_text = para.text
    # 找标签边界：匹配 "摘要：", "关键词：", "Abstract:", "Key words:" 等
    m = re.match(
        r'^(\s*(?:摘\s*要|关键\s*词|Abstract|Key\s*[Ww]ords?)\s*[:：]\s*)',
        full_text, re.IGNORECASE
    )
    if not m:
        # 找不到标签，全部用 content 格式
        for run in para.runs:
            _set_run_font(run, spec["content_font_cn"], spec["content_font_en"],
                          spec.get("content_size"), spec.get("content_bold", False))
        return

    label_len = len(m.group(1))
    consumed = 0
    for run in para.runs:
        run_len = len(run.text)
        if consumed + run_len <= label_len:
            # 标签部分
            _set_run_font(run, spec["label_font_cn"], spec["label_font_en"],
                          spec.get("label_size"), spec.get("label_bold", True))
        elif consumed >= label_len:
            # 内容部分
            _set_run_font(run, spec["content_font_cn"], spec["content_font_en"],
                          spec.get("content_size"), spec.get("content_bold", False))
        else:
            # 跨越边界的 run — 拆分不了，按内容处理
            _set_run_font(run, spec["content_font_cn"], spec["content_font_en"],
                          spec.get("content_size"), spec.get("content_bold", False))
        consumed += run_len


def _set_run_font(run, font_cn, font_en, size_pt=None, bold=False):
    """设置单个 run 的字体。"""
    from lxml import etree
    ns_w = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    font = run.font
    if font_en:
        font.name = font_en
    if size_pt:
        font.size = Pt(size_pt)
    font.bold = bold

    rpr = run._element.find(f"{ns_w}rPr")
    if rpr is None:
        rpr = etree.SubElement(run._element, f"{ns_w}rPr")
    if font_cn:
        fonts = rpr.find(f"{ns_w}rFonts")
        if fonts is None:
            fonts = etree.SubElement(rpr, f"{ns_w}rFonts")
        fonts.set(f"{ns_w}eastAsia", font_cn)


def apply_format(docx_path, rules, output_path=None):
    """将格式规则应用到文档。"""
    doc = Document(str(docx_path))

    # 页面设置
    apply_page_setup(doc, rules.get("_page", {}))

    # 跳过封面，从 TOC 之后开始应用
    seen_toc = False
    body_started = False
    stats = {k: 0 for k in rules if not k.startswith("_")}
    chapter_num = 0
    fig_count = 0
    table_count = 0

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        ptype = classify_paragraph(para)

        # 标记 TOC，但仍然应用格式
        if ptype == "toc":
            seen_toc = True
            # TOC 条目样式由 Word 自动管理，跳过
            continue
        if ptype == "toc_title":
            seen_toc = True
            # 继续应用 toc_title 格式

        # 跳过封面/声明
        if not seen_toc:
            continue

        # 跳过论文标题（大字号居中段落），保持原样
        if ptype == "title":
            body_started = True
            continue

        # 等到第一个标题出现才开始应用
        if not body_started:
            if ptype in ("heading1", "heading2", "heading3", "heading4", "ref_heading", "ack_heading", "toc_title"):
                body_started = True
            else:
                continue

        # 章节计数
        if ptype == "heading1":
            chapter_num += 1
            fig_count = 0
            table_count = 0

        # 图/表自动编号
        if ptype == "caption":
            new_text = renumber_caption(para.text.strip(), chapter_num,
                                        fig_count, table_count)
            if new_text != para.text.strip():
                update_caption_text(para, new_text)
            # 更新计数
            if para.text.strip().startswith("图") or para.text.strip().startswith("Fig"):
                fig_count += 1
            elif para.text.strip().startswith("表"):
                table_count += 1

        rule = rules.get(ptype)
        if not rule and ptype in ("heading1", "heading2", "heading3", "heading4"):
            rule = get_official_rule(ptype)
        elif rule and ptype in ("heading1", "heading2", "heading3", "heading4"):
            # 用官方规则覆盖关键字段
            official = get_official_rule(ptype)
            merged = {**rule, **official}
            rule = merged

        # 应用样式级格式（fallback）
        apply_style_fallback(para, rule)

        # 应用段落格式
        apply_para_format(para, rule)

        # 摘要/关键词：标签用黑体加粗，内容用楷体
        if ptype in ("abstract_cn", "keywords_cn"):
            apply_split_font(para, {
                "label_font_cn": "黑体", "label_font_en": "Times New Roman",
                "label_bold": True, "label_size": 10.5,
                "content_font_cn": "楷体", "content_font_en": "Times New Roman",
                "content_bold": False, "content_size": 10.5,
            })
        elif ptype in ("abstract_en", "keywords_en"):
            apply_split_font(para, {
                "label_font_cn": "宋体", "label_font_en": "Times New Roman",
                "label_bold": True, "label_size": 10.5,
                "content_font_cn": "宋体", "content_font_en": "Times New Roman",
                "content_bold": False, "content_size": 10.5,
            })
        # 标题：直接设置 run 级字体（从样式规则推导）
        elif ptype in ("heading1", "heading2", "heading3", "heading4"):
            font_cn = rule.get("style_font_cn") or rule.get("font_cn")
            font_en = rule.get("style_font_en") or rule.get("font_en", "Times New Roman")
            size_pt = rule.get("style_size_pt") or rule.get("size_pt")
            bold = rule.get("style_bold", True)
            for run in para.runs:
                _set_run_font(run, font_cn, font_en, size_pt, bold)
        else:
            # 应用字体格式到所有 runs
            for run in para.runs:
                apply_font(run, rule)

        # p<0.05 等统计标记斜体（正文和图注）
        if ptype in ("body", "caption"):
            italicize_statistics(para)

        stats[ptype] = stats.get(ptype, 0) + 1

    # 修复 TOC 样式（目录条目字体）
    fix_toc_styles(doc)

    # 设置 updateFields，打开 Word 时自动刷新目录和交叉引用
    update_fields(doc)

    # 检测引用错误
    ref_issues = check_references(doc)

    # 保存
    if not output_path:
        p = Path(docx_path)
        output_path = p.parent / f"{p.stem}_formatted{p.suffix}"
    doc.save(str(output_path))
    return output_path, stats, ref_issues


# ============================================================
# 检查差异
# ============================================================

def check_format(docx_path, rules):
    """检查文档格式与规则的差异。"""
    doc = Document(str(docx_path))
    issues = []

    seen_toc = False
    body_started = False

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        ptype = classify_paragraph(para)

        if ptype in ("toc", "toc_title"):
            seen_toc = True
            continue
        if not seen_toc:
            continue
        if not body_started:
            if ptype in ("heading1", "heading2", "heading3", "ref_heading", "ack_heading"):
                body_started = True
            else:
                continue

        rule = rules.get(ptype)
        if not rule:
            continue

        font = get_font_info(para)
        para_info = get_para_info(para)
        display_text = text[:40]

        # 检查字号
        if "size_pt" in rule and rule["size_pt"]:
            actual = font.get("size_pt")
            if actual and abs(actual - rule["size_pt"]) > 0.5:
                issues.append(f"[{ptype}] para {i}: 字号 {actual}pt != {rule['size_pt']}pt | {display_text}")

        # 检查加粗
        if "bold" in rule and rule["bold"] is not None:
            actual = font.get("bold", False)
            if actual != rule["bold"]:
                issues.append(f"[{ptype}] para {i}: bold={actual} != {rule['bold']} | {display_text}")

        # 检查对齐
        if "alignment" in rule and rule["alignment"]:
            actual = para_info.get("alignment", "")
            if actual and actual != rule["alignment"]:
                issues.append(f"[{ptype}] para {i}: align={actual} != {rule['alignment']} | {display_text}")

    return issues


# ============================================================
# CLI
# ============================================================

def main():
    if len(sys.argv) < 3:
        print(__doc__)
        sys.exit(1)

    cmd = sys.argv[1]

    if cmd == "learn":
        ref_path = sys.argv[2]
        output = sys.argv[3] if len(sys.argv) > 3 else "rules.json"
        rules = learn_format(ref_path)
        with open(output, "w", encoding="utf-8") as f:
            json.dump(rules, f, ensure_ascii=False, indent=2)
        print(f"Learned {len([k for k in rules if not k.startswith('_')])} format rules")
        print(f"Page: {rules.get('_page', {})}")
        for k, v in rules.items():
            if not k.startswith("_"):
                print(f"  {k}: {v}")
        print(f"Saved to: {output}")

    elif cmd == "apply":
        docx_path = sys.argv[2]
        rules_path = sys.argv[3]
        output_path = sys.argv[4] if len(sys.argv) > 4 else None
        with open(rules_path, "r", encoding="utf-8") as f:
            rules = json.load(f)
        output_path, stats, ref_issues = apply_format(docx_path, rules, output_path)
        print(f"Applied to: {output_path}")
        for k, count in stats.items():
            if count > 0:
                print(f"  {k}: {count} paragraphs")
        if ref_issues:
            print(f"\n⚠ 引用问题 ({len(ref_issues)}):")
            for issue in ref_issues:
                print(f"  {issue}")
        else:
            print("\n图/表引用检查通过")

    elif cmd == "check":
        docx_path = sys.argv[2]
        rules_path = sys.argv[3]
        with open(rules_path, "r", encoding="utf-8") as f:
            rules = json.load(f)
        issues = check_format(docx_path, rules)
        if issues:
            print(f"Found {len(issues)} issues:")
            for issue in issues:
                print(f"  {issue}")
        else:
            print("No issues found!")

    else:
        print(f"Unknown command: {cmd}")
        print(__doc__)
        sys.exit(1)


if __name__ == "__main__":
    main()
